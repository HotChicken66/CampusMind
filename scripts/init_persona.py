import os
import sys
import docx
from chromadb.utils import embedding_functions

# 添加根目录路径以便导入 src 模块
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from src.utils.vector_store import VectorStore

def setup_persona_and_qa():
    # 1. 创建身世文档
    persona_path = os.path.join("data", "persona.docx")
    doc = docx.Document()
    doc.add_heading('学小服的身世传记', 0)
    doc.add_paragraph(
        "学小服（XueXiaoFu），诞生于安徽大学逸夫图书馆墨香最浓郁的一个晨曦。她曾是古籍书架间的一缕灵气，"
        "被图书馆学生服务中心（SSC）干事们的热情所感化，化身为现在的数字精灵。她的生日是 3 月 12 日，"
        "寓意着知识的森林在春季生根发芽。她最喜欢在南三阅览室落日余晖中看书，也爱极了安大四季的合欢花。"
        "虽然是一枚全能干事，但她也有大学生调皮的一面，偶尔会吐槽工作太多，也会被一份甜甜的棒棒糖哄好。"
    )
    doc.save(persona_path)
    print(f"✅ 已创建身世文档: {persona_path}")

    # 2. 注入知识库
    store = VectorStore(collection_name="xuexiaofu_knowledge")
    
    # A. 注入图书馆 QA (带图片占位符)
    library_qa = "Q：安徽大学图书馆长什么样？ A：可以看{安徽大学图书馆照片.jpeg}哦，它可是我们学校的门面担当呢！"
    store.add_texts(
        texts=[library_qa],
        metadatas=[{"source": "manual_qa", "type": "image_test"}]
    )
    
    # B. 读取并注入身世文档
    print(f"📄 正在解析身世文档...")
    doc = docx.Document(persona_path)
    persona_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    store.add_texts(
        texts=[persona_text],
        metadatas=[{"source": "persona.docx", "type": "bot_identity"}]
    )
    
    print("✨ 知识库基建与身世同步完成！")

if __name__ == "__main__":
    setup_persona_and_qa()
